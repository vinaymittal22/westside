"""
Burnt Toast × TATA — AI Fashion Stylist
Generates full presentation document as both .docx and .pdf
"""

# ─────────────────────────────────────────────────────────────
# WORD DOCUMENT (.docx)
# ─────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# Brand Colors
CREAM   = RGBColor(0xF0, 0xEB, 0xE0)
INK     = RGBColor(0x1A, 0x1A, 0x1A)
ACCENT  = RGBColor(0xB8, 0x49, 0x2C)   # burnt orange
SAGE    = RGBColor(0x74, 0x8B, 0x6A)
GOLD    = RGBColor(0xC9, 0x96, 0x2E)
DARK_BG = RGBColor(0x1A, 0x1A, 0x1A)

def set_cell_bg(cell, hex_color):
    """Set background colour of a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'B8492C')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(22)
    run.font.color.rgb = ACCENT
    return p

def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(15)
    run.font.color.rgb = SAGE
    return p

def heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = INK
    return p

def body(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Cm(0.2)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(11)
    run.font.color.rgb = color if color else INK
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.6 + level * 0.4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = INK
    return p

def slide_divider(doc, slide_num, title):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run(f"── SLIDE {slide_num} ──")
    run.font.size  = Pt(9)
    run.font.color.rgb = GOLD
    run.bold = True
    heading1(doc, title)
    add_horizontal_rule(doc)

def add_comparison_table(doc):
    headers = ["Dimension", "Regular Chatbot", "Toastie (AI Stylist)"]
    rows = [
        ["Purpose",               "Customer service — answers FAQs",                     "Creative stylist — builds personalised looks"],
        ["Output",                "Text responses, order status, policy info",            "Full visual outfits with real products & prices"],
        ["Product Knowledge",     "Basic info (name, price, availability)",               "Deep aesthetic intelligence — color harmony, formality, occasion fit"],
        ["Personalisation",       "Rule-based, scripted",                                 "AI-generated, unique to every user prompt"],
        ["Fashion Intelligence",  "None",                                                 "8 aesthetics × 11 occasions × 10 product types"],
        ["User Experience",       "Transactional, functional",                            "Conversational, creative, emotionally engaging"],
        ["Revenue Impact",        "Near zero (deflects service calls)",                   "Direct — drives basket size, conversion, return visits"],
        ["Brand Value",           "Utility",                                              "Differentiation — a feature no competitor has"],
        ["Swapping Logic",        "Not applicable",                                       "Intelligent slot replacement with rejected-SKU memory"],
        ["Image Understanding",   "None",                                                 "Vision AI — recreates a look from an uploaded photo"],
    ]
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = 'Table Grid'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, "1A1A1A")
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = CREAM
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            if r_idx % 2 == 0:
                set_cell_bg(cell, "F7F3EE")
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            run.font.color.rgb = INK
            if c_idx == 0:
                run.bold = True
                run.font.color.rgb = ACCENT
    # Column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.6)
        row.cells[1].width = Inches(2.4)
        row.cells[2].width = Inches(2.9)

def add_metrics_table(doc):
    headers = ["Metric", "Current Baseline", "Toastie Target"]
    rows = [
        ["Avg. Order Value",   "₹1,200 – 1,500",     "₹2,500 – 3,500  (+80%)"],
        ["Session Duration",   "3 – 5 minutes",       "8 – 12 minutes  (+120%)"],
        ["Conversion Rate",    "Industry avg: ~2%",   "Target: 3 – 4%  (+50–100%)"],
        ["Repeat Visit Rate",  "Baseline",             "+40–60%  (AI personalization benchmark)"],
        ["Return Rate",        "Baseline",             "Reduced — confidence-driven purchasing"],
    ]
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, "B8492C")
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = CREAM
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            if r_idx % 2 == 0:
                set_cell_bg(cell, "F7F3EE")
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            run.font.color.rgb = INK
            if c_idx == 0:
                run.bold = True
            if c_idx == 2:
                run.font.color.rgb = SAGE
                run.bold = True
    for row in table.rows:
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(2.0)
        row.cells[2].width = Inches(2.9)

# ═══════════════════════════════════════════════════════════════
# BUILD THE DOCUMENT
# ═══════════════════════════════════════════════════════════════
doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── COVER PAGE ────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run("BURNT TOAST × TATA")
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = ACCENT

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("TOASTIE — India's First AI Fashion Stylist for Gen Z")
run2.bold = True
run2.font.size = Pt(18)
run2.font.color.rgb = INK

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("Turning Scroll into Style.  Style into Sales.")
run3.italic = True
run3.font.size = Pt(13)
run3.font.color.rgb = SAGE

doc.add_paragraph()
add_horizontal_rule(doc)
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Presented by: Anshul Mittal     |     Status: Demo & Testing Phase     |     May 2026").font.size = Pt(10)

doc.add_page_break()

# ── SLIDE 1 — AGENDA ──────────────────────────────────────────
slide_divider(doc, 1, "AGENDA — What We Will Cover Today")
body(doc, "This presentation walks you through six key areas:")
bullet(doc, "The Problem — Why Gen Z shoppers are stuck, and what it costs Burnt Toast")
bullet(doc, "The Solution — Introducing Toastie, the AI stylist built inside Burnt Toast")
bullet(doc, "How It Works — The full user experience, step by step")
bullet(doc, "Business Impact — Engagement, conversion, basket size, and revenue")
bullet(doc, "Why This Beats a Regular Chatbot — The difference that makes this special")
bullet(doc, "The Road Ahead — Where Toastie goes from here, and what we need to move forward")

doc.add_page_break()

# ── SLIDE 2 — THE PROBLEM ─────────────────────────────────────
slide_divider(doc, 2, "THE PROBLEM — The Hidden Cost of 'I Don't Know What to Wear'")

heading2(doc, "Who is Burnt Toast's Customer?")
body(doc, "Burnt Toast by Trent Limited is a youth-forward fashion brand launched in August 2025, "
     "specifically designed for India's Gen Z. The brand sits in a bold, intentional space — more "
     "trend-driven than Zudio, but more accessible than Zara or H&M. It speaks to young creators, "
     "trendsetters, and digital-first shoppers who don't just wear clothes — they express identity "
     "through what they wear.")
body(doc, "India is home to approximately 380 million Gen Z consumers — the largest youth cohort in "
     "the world. Their direct spending power in India is estimated at ₹20+ lakh crore ($250 billion) "
     "annually, with 47% of that going specifically to fashion and lifestyle. Myntra's Gen Z customer "
     "base alone doubled to 16 million users in a single year.", bold=False)

heading2(doc, "Problem 1 — Choice Overload & Decision Paralysis")
body(doc, "A typical Burnt Toast collection has hundreds of products spread across multiple categories. "
     "When a 19-year-old opens the site before a college fest, they are not looking for a catalogue — "
     "they are looking for a look. But the website gives them a grid. Thousands of items. No narrative. "
     "No direction. The result? They browse for 15–20 minutes, feel overwhelmed, and either leave "
     "without buying — or add one safe item they've seen before.")
body(doc, "This is called decision fatigue — and it is one of the biggest silent killers of online "
     "fashion conversion.", italic=True, color=ACCENT)

heading2(doc, "Problem 2 — No Personalisation, No Context")
body(doc, "Today's website treats everyone the same. The 18-year-old college student who wants a "
     "boho festival look gets the same homepage as the 24-year-old looking for a smart-casual "
     "brunch outfit. There is zero recognition of occasion, body type, style preference, or budget.")
body(doc, "Research confirms this is a massive lost opportunity: 50% of fashion purchases are now "
     "driven by personalisation. Brands that personalise see conversion rates increase by up to 40%. "
     "Brands that don't, watch their traffic bounce.")

heading2(doc, "Problem 3 — The Gap Between Discovery & Purchase")
body(doc, "Gen Z discovers fashion on Instagram, Pinterest, and YouTube — they see a full look in a "
     "Reel, get inspired, then land on an ecommerce site and cannot recreate that look. 76% of Gen Z "
     "fashion shoppers discover through social media. Only a fraction converts, because the path from "
     "'I love this vibe' to 'I found the exact outfit' is broken.")

heading3(doc, "What Does This Cost Burnt Toast?")
body(doc, "Every session that ends without a purchase is lost revenue. Every customer who buys only "
     "one item instead of a full look is a missed basket. Every shopper who cannot find what they're "
     "looking for becomes a Zara or H&M customer instead. Fixing this problem early is not optional — "
     "it is the difference between becoming India's defining Gen Z fashion brand, and becoming another "
     "label that Gen Z browses but doesn't commit to.")

doc.add_page_break()

# ── SLIDE 3 — THE SOLUTION ────────────────────────────────────
slide_divider(doc, 3, "THE SOLUTION — Meet Toastie, Your AI Fashion Stylist")

body(doc, "The answer is not more products. The answer is smarter discovery.", bold=True, color=ACCENT)
body(doc, "Toastie is an AI-powered fashion stylist — a conversational, visual, intelligent experience "
     "built directly inside the Burnt Toast website. It is not a search bar. It is not a filter. It is "
     "not a FAQ bot. It is a personal stylist, available 24/7, for every single customer who visits the site.")

heading2(doc, "What Toastie Does in Plain Language")
body(doc, "A user opens Burnt Toast's website. Instead of scrolling through hundreds of items alone, "
     "they simply talk to Toastie. They say: 'I have a college fest this weekend. I want something "
     "fun and Y2K-inspired. My budget is around ₹3,000.'")
body(doc, "In seconds, Toastie builds a complete, curated outfit — a top, bottom, footwear, bag, and "
     "jewelry — selected from Burnt Toast's real, live product catalogue. Every item is shoppable. "
     "Every price is real. The look is cohesive. The vibe is exactly what they asked for.")
body(doc, "Then they say: 'I love this, but can you show me different shoes? Something more grunge?' "
     "Toastie swaps just the footwear — keeping the rest of the outfit intact — and shows three "
     "alternative options. The user is now co-designing their own look in real time, with an AI "
     "that understands fashion.")

heading2(doc, "Three Things That Make Toastie Different")
heading3(doc, "1. It Speaks Gen Z's Language")
body(doc, "Toastie doesn't ask users to fill a form. It holds a natural conversation. Users say things "
     "like 'something cute for brunch,' 'give me a boho-coastal look,' or 'I want to look like I'm "
     "going to a rooftop party' — and Toastie understands all of it. It maps casual language to precise "
     "fashion aesthetics.")

heading3(doc, "2. It Builds Complete Looks, Not Just Recommendations")
body(doc, "This is the key distinction from any chatbot or recommendation engine. Toastie doesn't suggest "
     "one product at a time. It builds a full outfit — top, bottom, dress, outerwear, footwear, bag, "
     "jewelry, watch, hat — selected from real Burnt Toast inventory. It thinks the way a stylist thinks: "
     "color harmony, occasion fit, formality level, aesthetic coherence.")

heading3(doc, "3. It's Built on Real Products, Not Generic Advice")
body(doc, "Every outfit Toastie creates uses actual SKUs from Burnt Toast's live catalogue — with real "
     "prices, real images, real sizes, and a direct link to purchase. There is no gap between 'style "
     "advice' and 'add to cart.' The recommendation IS the product.")

doc.add_page_break()

# ── SLIDE 4 — HOW IT WORKS ────────────────────────────────────
slide_divider(doc, 4, "HOW IT WORKS — The Full User Experience, Step by Step")

steps = [
    ("Step 1 — The Conversation Begins",
     "The customer lands on Burnt Toast's website and sees Toastie's chat interface. The prompt is warm "
     "and Gen Z-native: 'Hey! What are we dressing for today? 🍞' The user responds in natural language "
     "— their occasion, their mood, their budget, their vibe. Even 'I don't know, just give me something "
     "cute' is enough to get started."),
    ("Step 2 — Toastie Reads the Brief",
     "Behind the scenes, Toastie uses Claude (Anthropic's AI model) to parse the user's message. It "
     "identifies: Occasion (college, date night, brunch, party, festival...), Aesthetic (Y2K revival, "
     "boho-coastal, urban streetwear...), Budget, Gender, and Preferred colors. This is not keyword "
     "matching — this is genuine language understanding."),
    ("Step 3 — The Outfit is Built",
     "Toastie's outfit engine runs across the full Burnt Toast catalogue — 280+ curated products. It "
     "selects the right top or dress, a complementary bottom, matching footwear, a curated bag, and "
     "jewelry that ties the look together. Every selection is made based on aesthetic alignment, color "
     "harmony, formality level, boldness balance, and budget fit."),
    ("Step 4 — The Look is Revealed",
     "The customer sees their complete outfit — displayed as a lookbook-style card with real product "
     "images, product names and prices, total outfit cost, a personality-driven description written "
     "in Burnt Toast's brand voice, and a vibe label like 'Sunday Brunch Slay' or 'Main Character "
     "Energy.' Every slot is shoppable — one click goes to the product page."),
    ("Step 5 — The Customer Customises",
     "The customer is not locked in. They can swap any single piece, filter by subtype ('Show me only "
     "sneakers'), change the whole vibe, lock what they love, set a budget, or add/remove items. Every "
     "change is handled conversationally. They are not navigating menus — they are talking to their stylist."),
    ("Step 6 — Add to Cart & Purchase",
     "When the customer loves their look, they click on any product to go directly to its product page. "
     "The purchase journey is standard Burnt Toast — familiar, trusted, fast. The difference: they arrive "
     "with intention. They know exactly what they want, and they want all of it."),
]
for title, text in steps:
    heading3(doc, title)
    body(doc, text)

doc.add_page_break()

# ── SLIDE 5 — KEY FEATURES ────────────────────────────────────
slide_divider(doc, 5, "KEY FEATURES — Full Capability List")

features = [
    ("Feature 1 — Complete Outfit Generation",
     "Generates a full, styled outfit from scratch based on natural language. Covers every slot: "
     "tops, bottoms, dresses, outerwear, footwear, bags, jewelry (necklaces, earrings, bracelets, "
     "rings, charms), watches, eyewear, and hats. All jewelry shown under a unified JEWELLERY label."),
    ("Feature 2 — Aesthetic Intelligence (8 Style Profiles)",
     "Y2K Revival · Urban Streetwear · Smart Casual · Minimal Clean · Boho Coastal · "
     "Preppy Collegiate · Athleisure · Feminine Romantic — each with its own logic for "
     "silhouette, color palette, fabric, and accessory style."),
    ("Feature 3 — Occasion Mapping (11 Occasions)",
     "Casual · College · Brunch · Date Night · Party · Festival · Beach · Travel · Work · "
     "Active · Wedding — Toastie knows what each occasion demands in terms of formality, "
     "vibe, and appropriateness."),
    ("Feature 4 — Real-Time Slot Swapping",
     "Any individual piece can be swapped without disrupting the rest of the outfit. Works for: "
     "top, bottom, dress, footwear, bag, jewelry, outerwear, and accessories."),
    ("Feature 5 — Footwear Subtype Filter",
     "When a customer explicitly asks for a specific footwear type — sneakers, sandals, loafers, "
     "heels, boots, ballerinas, mary janes, mules, platforms — Toastie shows only that type. "
     "No mixing. No irrelevant suggestions."),
    ("Feature 6 — Multi-Outfit Variations",
     "Generates 3 complete outfit options at once, each varying on a specific slot (e.g., three "
     "different tops with the same bottom and shoes)."),
    ("Feature 7 — Budget Intelligence",
     "Works within price constraints across all slots combined. Provides a budget note explaining "
     "how the full look fits within the user's stated spend."),
    ("Feature 8 — Product Memory & Rejection",
     "Within a session, Toastie remembers which products the user has rejected. It never "
     "re-suggests a dismissed item, even across multiple swaps."),
    ("Feature 9 — Claude Vision — Image-Based Inspiration",
     "Users can upload a photo — a screenshot from Instagram, Pinterest, or a street style shot — "
     "and Toastie will analyse it and recreate a similar look using Burnt Toast's real catalogue."),
    ("Feature 10 — Brand-Voice Copy Generation",
     "Every outfit includes a hype caption written in Burnt Toast's voice. Not generic. Not "
     "corporate. Something like: 'Coastal cool meets Sunday ease. You're not just going to "
     "brunch — you're arriving.'"),
]
for title, text in features:
    heading3(doc, title)
    body(doc, text)

doc.add_page_break()

# ── SLIDE 6 — GEN Z OPPORTUNITY ───────────────────────────────
slide_divider(doc, 6, "THE GEN Z OPPORTUNITY — Why Now. Why This Audience.")

heading2(doc, "India's Gen Z — The Numbers That Matter")
stats = [
    "380 million Gen Z consumers in India — the largest Gen Z population of any country in the world",
    "$250 billion in direct annual spending; 47% goes to fashion and lifestyle",
    "By 2035, Gen Z projected to drive $2 trillion in total consumer spending in India",
    "44% of all smartphone purchases in India in 2024 came from Gen Z",
    "Myntra's Gen Z base doubled to 16 million users in a single year",
]
for s in stats:
    bullet(doc, s)

heading2(doc, "How Gen Z Shops — And Why Traditional Retail Fails Them")
body(doc, "Gen Z does not shop the way their parents did. They trust experiences, aesthetics, and "
     "community over ads and static product grids. Here is what research tells us:")
stats2 = [
    "76% discover fashion through social media — primarily Instagram, YouTube Shorts, and Pinterest",
    "39% of Indian Gen Z consumers buy fashion after watching Instagram Reels",
    "75% of Gen Z digital commerce happens on smartphones",
    "50% of fashion purchases are driven by personalisation",
    "61% actively used AI tools to help with a purchase in the past year",
    "75% are interested in using AI during the shopping process — highest of any generation",
]
for s in stats2:
    bullet(doc, s)

heading2(doc, "The Competitive Window")
body(doc, "Right now, no Indian fashion brand at Burnt Toast's price point and positioning has an AI "
     "stylist. Not Zudio. Not H&M India. Not Bershka. Not Stradivarius. This is a first-mover window. "
     "The brand that builds the most trusted AI styling relationship with Indian Gen Z first will capture "
     "loyalty that is extremely hard to win back later. Toastie gives Burnt Toast the chance to be that brand.",
     color=ACCENT)

doc.add_page_break()

# ── SLIDE 7 — BUSINESS IMPACT: ENGAGEMENT ─────────────────────
slide_divider(doc, 7, "BUSINESS IMPACT — Engagement")

heading2(doc, "Metric 1 — Time on Site")
body(doc, "Today, a typical fashion ecommerce session in India lasts 3–5 minutes. A Toastie session is "
     "a conversation — it has narrative momentum. AI-fashion integrations globally show conversational "
     "styling experiences increase average session time by 35–60%. Every additional minute a customer "
     "spends on your platform increases the statistical probability of a purchase.")

heading2(doc, "Metric 2 — Return Visits & Retention")
body(doc, "Toastie creates a reason to return that no static catalogue can provide. A customer who used "
     "Toastie for their college fest outfit will think of it again for their next occasion. Brands using "
     "AI personalisation report 40–60% improvement in repeat visit rates among engaged users vs. those "
     "who browsed without the tool.")

heading2(doc, "Metric 3 — Depth of Catalogue Engagement")
body(doc, "Without Toastie, customers typically discover 10–20 products per session. With Toastie, the "
     "outfit engine intelligently surfaces products from across the entire 280+ item catalogue based on "
     "aesthetic, occasion, and coherence — not just what's trending on the homepage. Hidden inventory "
     "becomes discovered opportunity.")

heading2(doc, "Metric 4 — Net Promoter Effect")
body(doc, "When Toastie gives someone a look they love, they screenshot it, share it on Instagram Stories, "
     "and tag Burnt Toast. The AI stylist becomes a conversation starter — 'I just built this outfit on "
     "Burnt Toast's website' — which is organic, user-generated marketing that no ad spend can buy.")

doc.add_page_break()

# ── SLIDE 8 — BUSINESS IMPACT: CONVERSION & REVENUE ──────────
slide_divider(doc, 8, "BUSINESS IMPACT — Conversion & Revenue")

heading2(doc, "Impact 1 — Higher Conversion Rate")
body(doc, "The global average conversion rate for fashion ecommerce is 1.5–3%. The fundamental reason "
     "is intent mismatch — most visitors are browsing, not buying, because they haven't yet found "
     "something they're confident about. Toastie resolves the confidence problem.")
body(doc, "AI-powered personalisation in fashion ecommerce has been shown to increase conversion rates "
     "by 30–40%. Visual search and AI recommendation tools show 27% higher conversion in controlled "
     "studies. Even a 25% lift in conversion rate at Burnt Toast's stage is transformational.", bold=True, color=SAGE)

heading2(doc, "Impact 2 — Larger Average Basket Size")
body(doc, "Without Toastie: Customer finds a top they like, adds to cart. Average transaction: ₹1,200–1,500.")
body(doc, "With Toastie: Customer builds a complete outfit — top (₹799) + bottom (₹999) + footwear (₹1,299) "
     "+ bag (₹899) + jewelry (₹399). Even if they buy 3 of 5 pieces, the transaction is ₹2,500–3,500.")
body(doc, "Kering's AI styling app (Luce) reported 15–20% increase in average order value. Global AI chatbot "
     "implementations report basket size increases of 20–35% when customers engage with outfit recommendations.")
body(doc, "For Burnt Toast, every ₹500 increase in average order value across 1,000 transactions/month = "
     "₹6 lakh in additional monthly revenue — with zero additional customer acquisition cost.", bold=True, color=ACCENT)

heading2(doc, "Impact 3 — Reduced Return Rates")
body(doc, "When customers buy an item that doesn't go with what they own, or doesn't fit the occasion "
     "they imagined, they return it. Toastie reduces this by design. When a customer builds a complete "
     "look and purchases multiple pieces styled together, they buy with confidence and context. Lower "
     "return rates mean lower operational cost, lower logistics cost, and higher net margin per order.")

heading2(doc, "Impact 4 — New Customer Acquisition Through Experience")
body(doc, "For Burnt Toast, competing against Zudio on price or against H&M on brand recognition is "
     "hard. But competing on experience — offering something no competitor has — is where Burnt Toast "
     "wins. Toastie is that experience. It is a reason to choose Burnt Toast specifically.")
body(doc, "In the Gen Z market, experience IS the product.", bold=True, italic=True, color=ACCENT)

doc.add_page_break()

# ── SLIDE 9 — TOASTIE VS CHATBOT ─────────────────────────────
slide_divider(doc, 9, "TOASTIE vs. A REGULAR CHATBOT — Why This Is Not a Bot")

body(doc, "A standard chatbot is a customer service tool. It answers questions like 'Where is my "
     "order?' and 'What is your return policy?' It is reactive. It is scripted. It has no understanding "
     "of aesthetics, no awareness of your product catalogue's style logic, and no ability to build "
     "something original. Toastie is a generative, creative AI system. The difference is fundamental.")

doc.add_paragraph()
add_comparison_table(doc)
doc.add_paragraph()

body(doc, "In Short:", bold=True, color=ACCENT)
bullet(doc, "A chatbot saves time.  Toastie creates desire.  And in fashion, desire drives purchase.")
bullet(doc, "A chatbot is a cost reduction tool.  Toastie is a revenue generation engine.")
bullet(doc, "A chatbot is invisible when it works.  Toastie becomes a brand asset — a feature customers come back for.")

doc.add_page_break()

# ── SLIDE 10 — FUTURE ROADMAP ─────────────────────────────────
slide_divider(doc, 10, "FUTURE ROADMAP — 4 Phases of Growth")

phases = [
    ("PHASE 1 — APPROVE & REFINE  (Current: Q2–Q3 2026)",
     "ACCENT",
     ["Core AI stylist built, tested, and functional",
      "Full product catalogue (280+ products) integrated",
      "All key features live — outfit generation, swap logic, footwear filter, vision input, budget intelligence",
      "Goal: panel approval, technical validation, user testing with real Burnt Toast customers"]),
    ("PHASE 2 — LAUNCH & LEARN  (Q4 2026 – Q1 2027)",
     "SAGE",
     ["Toastie goes live on the Burnt Toast website for all users",
      "Performance measurement — conversion, basket size, session time, return rate",
      "New features: wishlist integration, 'Shop the Look' shareable pages, WhatsApp/Instagram DM integration",
      "User feedback loop: 'Rate this look', thumbs up/down on suggestions"]),
    ("PHASE 3 — SCALE & PERSONALISE  (Q2–Q3 2027)",
     "GOLD",
     ["User profiles — Toastie remembers your style across sessions",
      "Style evolution tracking — notices when your aesthetic shifts and adapts",
      "Size intelligence — remembers your size preferences across all product types",
      "In-store kiosk mode — Toastie on a screen inside Burnt Toast physical stores"]),
    ("PHASE 4 — MONETISE & EXPAND  (2028+)",
     "INK",
     ["Trend prediction engine — Toastie's user data creates a proprietary trend signal for Burnt Toast's buying team",
      "Creator collaboration — influencers build and share Burnt Toast looks via Toastie",
      "Cross-brand expansion — engine extended to Zudio, Westside, or Tata CLiQ",
      "White-label licensing — AI stylist technology offered to other fashion brands as a platform play"]),
]
color_map = {"ACCENT": ACCENT, "SAGE": SAGE, "GOLD": GOLD, "INK": INK}
for title, clr, items in phases:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = color_map[clr]
    for item in items:
        bullet(doc, item)

doc.add_page_break()

# ── SLIDE 11 — MARKET OPPORTUNITY ────────────────────────────
slide_divider(doc, 11, "MARKET OPPORTUNITY — The Size of the Prize")

heading2(doc, "The Indian Fashion Market — Today")
stats3 = [
    "India clothing & apparel market valued at ₹9.7 lakh crore ($116.64 billion) in 2025",
    "Growing at 4.3% CAGR — projected to reach ₹14.8 lakh crore ($177.70 billion) by 2035",
    "Online fashion ecommerce: $21.6 billion in 2025 → $98.45 billion by 2032 (24.2% CAGR)",
    "Fashion ecommerce share: 15% of sales in 2023 → forecasted 25% by 2030",
]
for s in stats3:
    bullet(doc, s)

heading2(doc, "The Gen Z Fashion Segment")
stats4 = [
    "380 million Gen Z consumers in India",
    "$250 billion in direct annual spending, 47% on fashion/lifestyle",
    "Projected to drive $2 trillion in total spending by 2035",
    "Myntra's Gen Z base: 8 million → 16 million in a single year",
]
for s in stats4:
    bullet(doc, s)

heading2(doc, "The AI Fashion Commerce Market — Global Context")
stats5 = [
    "Generative AI in fashion market expected to grow to $4.4 billion by 2028",
    "97% of commerce organisations now have AI plans in place",
    "75% of Gen Z shoppers interested in using AI during the purchase process",
    "AI-powered experiences deliver 30–40% conversion lift and 15–25% basket size increase",
]
for s in stats5:
    bullet(doc, s)

heading2(doc, "The Trent Advantage")
body(doc, "Burnt Toast does not enter this market from zero. It enters with:")
stats6 = [
    "Tata Group's trust, financial strength, and retail infrastructure",
    "Trent FY25: ₹17,353 Cr revenue (+37%), 1,043 stores across 242 cities",
    "Zudio: 765 stores in 235 cities, crossed $1 billion in sales in FY25",
    "Physical stores in Bengaluru, Pune, Surat, and Thane ready for Toastie-enabled rollout",
]
for s in stats6:
    bullet(doc, s)

doc.add_page_break()

# ── SLIDE 12 — CUSTOMER EXPERIENCE JOURNEY ───────────────────
slide_divider(doc, 12, "CUSTOMER EXPERIENCE JOURNEY — Before vs. After Toastie")

body(doc, "Meet Priya. She's 20. She's in Mumbai. She has a college farewell party in three days.", bold=True)

heading2(doc, "BEFORE TOASTIE — Without the AI Stylist")
body(doc, "Priya opens Burnt Toast's website. She types 'dresses' in the search bar. 47 results appear. "
     "She scrolls. Some are too formal. Some are too casual. Some are the right vibe but she's not sure "
     "what shoes would go with them. She opens three products in new tabs. She reads descriptions. She "
     "closes two. She adds one to her cart. Then she goes to check Instagram. She sees a Reel of a girl "
     "in a perfect Y2K look. She gets inspired again. She goes back to the site. She can't find what "
     "she saw. She leaves. She buys something from Myntra instead.")
body(doc, "Burnt Toast lost a customer. Not because the product wasn't there. Because the discovery "
     "experience failed her.", bold=True, color=ACCENT)

heading2(doc, "AFTER TOASTIE — With the AI Stylist")
body(doc, "Priya opens Burnt Toast's website. Toastie's chat appears: 'Hey! What are we dressing for? 🍞'")
body(doc, "She types: 'College farewell. Want to look iconic. Y2K-ish. Budget around ₹3,500.'")
body(doc, "In 4 seconds, Toastie shows her a complete look:")
items = [
    "Rust-orange co-ord set — ₹1,299",
    "Platform sneakers in white — ₹1,199",
    "Mini chain bag — ₹599",
    "Hoop earrings + chunky bracelet — ₹399",
    "Total: ₹3,496 — within budget ✓",
]
for item in items:
    bullet(doc, item)
body(doc, "The look is labelled: 'Farewell Queen Energy.'")
body(doc, "Caption: 'Low-rise, high vibe. You're not leaving college — you're leaving a legacy.'", italic=True)
body(doc, "Priya loves the co-ord but wants different shoes. She types: 'Can I see some heels instead?' "
     "Toastie swaps just the footwear — shows three heel options within budget. She picks the strappy "
     "kitten heels. She screenshots the look. She adds all 5 items to her cart.")
body(doc, "₹3,496 average order. 5 products. 8-minute session. Zero ad spend.", bold=True, color=SAGE)
body(doc, "Burnt Toast didn't just make a sale. It made a loyal customer.", bold=True, color=ACCENT)

doc.add_page_break()

# ── SLIDE 13 — DEVELOPMENT STATUS ────────────────────────────
slide_divider(doc, 13, "DEVELOPMENT STATUS — What's Built, What's Next, What We Need")

heading2(doc, "What Has Been Built (Demo & Testing Phase — Fully Functional Now)")
built = [
    "✅  Conversational AI interface — Natural language input, Claude-powered intent parsing, full chat UI in Burnt Toast brand design",
    "✅  Outfit generation engine — Selects from 280+ curated products across all categories",
    "✅  8 aesthetic profiles — Y2K Revival, Urban Streetwear, Smart Casual, Minimal Clean, Boho Coastal, Preppy Collegiate, Athleisure, Feminine Romantic",
    "✅  11 occasion types — Casual, College, Brunch, Date Night, Party, Festival, Beach, Travel, Work, Active, Wedding",
    "✅  Real-time slot swapping — Any piece in any outfit can be individually swapped",
    "✅  Footwear subtype filter — Sneakers, sandals, loafers, heels, boots, ballerinas, mary janes, mules, platforms",
    "✅  Budget intelligence — Works within total price constraints across all slots",
    "✅  Multi-outfit variations — Generates 3 full look options varying on one slot",
    "✅  Product memory — Never re-shows rejected SKUs within a session",
    "✅  Vision AI — Upload a photo and recreate the look from Burnt Toast's catalogue",
    "✅  Claude-generated copy — Every outfit gets a branded, Gen Z-native caption",
    "✅  Full product catalogue — 280+ products with proper category mapping, pricing, images, aesthetics, and occasion tags",
]
for item in built:
    bullet(doc, item)

heading2(doc, "Technology Stack")
stack = [
    "Frontend: Next.js 16, React 19, TypeScript — fast, scalable, production-ready",
    "AI Layer: Anthropic Claude (Sonnet) — for language understanding and image vision",
    "Outfit Engine: Custom TypeScript logic — built and maintained in-house",
    "Product Data: Shopify-compatible product feed — ready for live catalogue sync",
    "Deployment: Vercel-ready — can go live within days of approval",
]
for s in stack:
    bullet(doc, s)

heading2(doc, "What We Need to Proceed")
needs = [
    "Panel Approval — Sign-off from the TATA business team to integrate Toastie into the live Burnt Toast website",
    "Live Product Feed Access — API or data feed from Burnt Toast's Shopify backend for real-time inventory sync",
    "Phase 2 Budget — Development resources for 3 months: analytics, wishlist, social sharing, and user feedback loop",
]
for n in needs:
    bullet(doc, n)
body(doc, "Timeline from approval to live: 6–8 weeks.", bold=True, color=ACCENT)

doc.add_page_break()

# ── SLIDE 14 — CLOSING: THE ASK ───────────────────────────────
slide_divider(doc, 14, "THE ASK — What We Are Asking the TATA Panel For")

body(doc, "What Toastie Is — One Line:", bold=True, color=ACCENT)
body(doc, "An AI-powered personal stylist built inside Burnt Toast that turns a shopper's vibe into a "
     "complete, shoppable look in seconds — driving higher conversion, larger baskets, and deeper brand "
     "loyalty among Gen Z.", italic=True)

heading2(doc, "Where We Stand")
status = [
    "✅  Product is built and functional",
    "✅  Catalogue is integrated (280+ products)",
    "✅  All core features are tested and working",
    "✅  Technology stack is production-ready",
    "✅  Zero external dependencies to go live",
]
for s in status:
    bullet(doc, s)
body(doc, "All we need is the green light.", bold=True, color=ACCENT)

heading2(doc, "Our Three Specific Asks")
asks = [
    ("Ask 1 — Approval to integrate Toastie into the live Burnt Toast website",
     "Grant us permission to connect Toastie's engine to the live burnt-toast.com product feed "
     "and launch a beta version for real users — even if limited to a test cohort initially."),
    ("Ask 2 — Access to live Burnt Toast product data",
     "To keep Toastie's recommendations accurate and in-stock, we need access to the live Shopify "
     "product feed. Toastie only reads product information already publicly visible on the website."),
    ("Ask 3 — A Phase 2 development budget",
     "To build the features that take Toastie from functional to exceptional — wishlist saving, "
     "social sharing, analytics dashboard, A/B testing — we estimate a 3-month runway of "
     "engineering resources."),
]
for title, text in asks:
    heading3(doc, title)
    body(doc, text)

heading2(doc, "The Opportunity in Numbers")
doc.add_paragraph()
add_metrics_table(doc)
doc.add_paragraph()

heading2(doc, "The Closing Thought")
body(doc, "Burnt Toast was built to be different. Not just another affordable fashion brand — but India's "
     "Gen Z fashion identity. Toastie is what makes that difference tangible, digital, and scalable.")
body(doc, "Every other feature on a fashion website — better photography, faster shipping, easier returns "
     "— can be copied. An AI stylist that knows your vibe, builds your look, and speaks your language "
     "is a relationship. And relationships, once built, are remarkably hard to lose.")
body(doc, "Toastie is not a feature. It is Burnt Toast's competitive moat.", bold=True, color=ACCENT)
body(doc, "We are ready to build it. We just need you to say yes.", bold=True, italic=True)

doc.add_paragraph()
add_horizontal_rule(doc)
p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_end = p_end.add_run('"Turning Scroll into Style.  Style into Sales."')
run_end.italic = True
run_end.font.size = Pt(13)
run_end.font.color.rgb = ACCENT

p_brand = doc.add_paragraph()
p_brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_brand.add_run("— Toastie, by Burnt Toast × TATA").font.size = Pt(11)

doc.add_paragraph()
p_conf = doc.add_paragraph()
p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_conf = p_conf.add_run("This document is confidential and prepared for internal TATA panel review only. May 2026.")
run_conf.font.size = Pt(9)
run_conf.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── SAVE .docx ────────────────────────────────────────────────
docx_path = "Burnt_Toast_Presentation.docx"
doc.save(docx_path)
print(f"✅  Saved: {docx_path}")


# ═══════════════════════════════════════════════════════════════
# PDF GENERATION
# ═══════════════════════════════════════════════════════════════
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm, mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                 Table, TableStyle, HRFlowable,
                                 PageBreak, KeepTogether)
from reportlab.platypus.flowables import HRFlowable

# --- Colours ---
C_ACCENT  = colors.HexColor("#B8492C")
C_SAGE    = colors.HexColor("#748B6A")
C_GOLD    = colors.HexColor("#C9962E")
C_INK     = colors.HexColor("#1A1A1A")
C_CREAM   = colors.HexColor("#F0EBE0")
C_LIGHT   = colors.HexColor("#F7F3EE")
C_GREY    = colors.HexColor("#999999")
C_WHITE   = colors.white
C_DARK    = colors.HexColor("#1A1A1A")

W, H = A4  # 595 x 842 pt
MARGIN = 2.2 * cm

# --- Styles ---
styles = getSampleStyleSheet()

def S(name, **kw):
    base = kw.pop("parent", "Normal")
    s = ParagraphStyle(name, parent=styles[base], **kw)
    return s

sCover1  = S("Cover1",   fontSize=28, textColor=C_ACCENT,  leading=34, alignment=TA_CENTER, fontName="Helvetica-Bold", spaceAfter=8)
sCover2  = S("Cover2",   fontSize=16, textColor=C_INK,     leading=22, alignment=TA_CENTER, fontName="Helvetica-Bold", spaceAfter=6)
sCover3  = S("Cover3",   fontSize=12, textColor=C_SAGE,    leading=18, alignment=TA_CENTER, fontName="Helvetica-Oblique", spaceAfter=4)
sMeta   = S("Meta",     fontSize=9,  textColor=C_GREY,    leading=14, alignment=TA_CENTER, spaceAfter=4)

sSlideNum = S("SlideNum", fontSize=8,  textColor=C_GOLD, fontName="Helvetica-Bold", spaceAfter=2, spaceBefore=18)
sH1     = S("H1",       fontSize=18, textColor=C_ACCENT, fontName="Helvetica-Bold", leading=24, spaceBefore=12, spaceAfter=4)
sH2     = S("H2",       fontSize=13, textColor=C_SAGE,   fontName="Helvetica-Bold", leading=18, spaceBefore=10, spaceAfter=3)
sH3     = S("H3",       fontSize=11, textColor=C_INK,    fontName="Helvetica-Bold", leading=16, spaceBefore=8,  spaceAfter=2)
sBody   = S("Body",     fontSize=10, textColor=C_INK,    leading=16, spaceBefore=3, spaceAfter=5, alignment=TA_JUSTIFY)
sBold   = S("Bold",     fontSize=10, textColor=C_INK,    leading=16, fontName="Helvetica-Bold", spaceBefore=3, spaceAfter=5)
sAccent = S("Accent",   fontSize=10, textColor=C_ACCENT, leading=16, fontName="Helvetica-Bold", spaceBefore=3, spaceAfter=5)
sSage   = S("Sage",     fontSize=10, textColor=C_SAGE,   leading=16, fontName="Helvetica-Bold", spaceBefore=3, spaceAfter=5)
sItalic = S("Italic",   fontSize=10, textColor=C_INK,    leading=16, fontName="Helvetica-Oblique", spaceBefore=3, spaceAfter=5)
sBullet = S("Bullet",   fontSize=10, textColor=C_INK,    leading=15, leftIndent=14, firstLineIndent=0,
             spaceBefore=2, spaceAfter=2,
             bulletIndent=4, bulletFontName="Helvetica", bulletFontSize=10)

def HR():
    return HRFlowable(width="100%", thickness=1.5, color=C_ACCENT, spaceAfter=6, spaceBefore=6)

def slide_header_pdf(num, title):
    return [
        Paragraph(f"── SLIDE {num} ──", sSlideNum),
        Paragraph(title, sH1),
        HR(),
        Spacer(1, 4),
    ]

def bul(text):
    return Paragraph(f"• {text}", sBullet)

def body_p(text, style=None):
    return Paragraph(text, style or sBody)

# ── comparison table ──
def comparison_table_pdf():
    data = [
        ["Dimension", "Regular Chatbot", "Toastie (AI Stylist)"],
        ["Purpose",             "Customer service — FAQs",            "Creative stylist — builds looks"],
        ["Output",              "Text / order status / policy",        "Full outfits with real products & prices"],
        ["Product Knowledge",   "Basic info",                          "Deep aesthetic intelligence"],
        ["Personalisation",     "Rule-based, scripted",                "AI-generated, unique every time"],
        ["Fashion Intelligence","None",                                 "8 aesthetics × 11 occasions × 10 types"],
        ["User Experience",     "Transactional",                       "Conversational, creative, engaging"],
        ["Revenue Impact",      "Near zero",                           "Drives basket size, conversion, loyalty"],
        ["Brand Value",         "Utility",                             "Differentiation — no competitor has it"],
        ["Swapping Logic",      "Not applicable",                      "Intelligent slot replacement + SKU memory"],
        ["Image Understanding", "None",                                "Vision AI — recreate look from a photo"],
    ]
    col_w = [(W - 2*MARGIN) * x for x in [0.24, 0.33, 0.43]]
    ts = TableStyle([
        ("BACKGROUND",  (0,0), (-1,0),  C_DARK),
        ("TEXTCOLOR",   (0,0), (-1,0),  C_CREAM),
        ("FONTNAME",    (0,0), (-1,0),  "Helvetica-Bold"),
        ("FONTSIZE",    (0,0), (-1,-1), 8.5),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [C_LIGHT, C_WHITE]),
        ("TEXTCOLOR",   (0,1), (0,-1),  C_ACCENT),
        ("FONTNAME",    (0,1), (0,-1),  "Helvetica-Bold"),
        ("GRID",        (0,0), (-1,-1), 0.4, C_GREY),
        ("VALIGN",      (0,0), (-1,-1), "MIDDLE"),
        ("TOPPADDING",  (0,0), (-1,-1), 4),
        ("BOTTOMPADDING",(0,0),(-1,-1), 4),
        ("LEFTPADDING", (0,0), (-1,-1), 5),
    ])
    t = Table(data, colWidths=col_w)
    t.setStyle(ts)
    return t

def metrics_table_pdf():
    data = [
        ["Metric", "Current Baseline", "Toastie Target"],
        ["Avg. Order Value",   "₹1,200 – 1,500",      "₹2,500 – 3,500  (+80%)"],
        ["Session Duration",   "3 – 5 minutes",        "8 – 12 minutes  (+120%)"],
        ["Conversion Rate",    "Industry avg: ~2%",    "Target: 3 – 4%  (+50–100%)"],
        ["Repeat Visit Rate",  "Baseline",              "+40–60%  (AI personalisation benchmark)"],
        ["Return Rate",        "Baseline",              "Reduced — confidence-driven purchasing"],
    ]
    col_w = [(W - 2*MARGIN) * x for x in [0.30, 0.32, 0.38]]
    ts = TableStyle([
        ("BACKGROUND",  (0,0), (-1,0),  C_ACCENT),
        ("TEXTCOLOR",   (0,0), (-1,0),  C_WHITE),
        ("FONTNAME",    (0,0), (-1,0),  "Helvetica-Bold"),
        ("FONTSIZE",    (0,0), (-1,-1), 9),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [C_LIGHT, C_WHITE]),
        ("FONTNAME",    (0,1), (0,-1),  "Helvetica-Bold"),
        ("TEXTCOLOR",   (2,1), (2,-1),  C_SAGE),
        ("FONTNAME",    (2,1), (2,-1),  "Helvetica-Bold"),
        ("GRID",        (0,0), (-1,-1), 0.4, C_GREY),
        ("VALIGN",      (0,0), (-1,-1), "MIDDLE"),
        ("TOPPADDING",  (0,0), (-1,-1), 5),
        ("BOTTOMPADDING",(0,0),(-1,-1), 5),
        ("LEFTPADDING", (0,0), (-1,-1), 6),
    ])
    t = Table(data, colWidths=col_w)
    t.setStyle(ts)
    return t

# ── Build story ──
pdf_path = "Burnt_Toast_Presentation.pdf"
doc_pdf = SimpleDocTemplate(pdf_path, pagesize=A4,
    leftMargin=MARGIN, rightMargin=MARGIN,
    topMargin=MARGIN, bottomMargin=MARGIN)

story = []

# Cover
story += [
    Spacer(1, 3*cm),
    Paragraph("BURNT TOAST × TATA", sCover1),
    Paragraph("TOASTIE — India's First AI Fashion Stylist for Gen Z", scover2 := S("c2", fontSize=15, textColor=C_INK, fontName="Helvetica-Bold", alignment=TA_CENTER, spaceAfter=8)),
    Paragraph("Turning Scroll into Style.  Style into Sales.", scover3 := S("c3", fontSize=12, textColor=C_SAGE, fontName="Helvetica-Oblique", alignment=TA_CENTER, spaceAfter=6)),
    Spacer(1, 0.5*cm),
    HR(),
    Spacer(1, 0.3*cm),
    Paragraph("Presented by: Anshul Mittal  |  Status: Demo & Testing Phase  |  May 2026", sMeta),
    PageBreak(),
]

# Slide 1 — Agenda
story += slide_header_pdf(1, "AGENDA — What We Will Cover Today")
story += [body_p("This presentation walks you through six key areas:")]
for item in [
    "The Problem — Why Gen Z shoppers are stuck, and what it costs Burnt Toast",
    "The Solution — Introducing Toastie, the AI stylist built inside Burnt Toast",
    "How It Works — The full user experience, step by step",
    "Business Impact — Engagement, conversion, basket size, and revenue",
    "Why This Beats a Regular Chatbot — The difference that makes this special",
    "The Road Ahead — Where Toastie goes from here, and what we need to move forward",
]:
    story.append(bul(item))
story.append(PageBreak())

# Slide 2 — The Problem
story += slide_header_pdf(2, "THE PROBLEM — The Hidden Cost of 'I Don't Know What to Wear'")
story += [
    Paragraph("Who is Burnt Toast's Customer?", sH2),
    body_p("Burnt Toast by Trent Limited is a youth-forward fashion brand launched in August 2025, "
           "specifically designed for India's Gen Z. It sits between Zudio and Zara — more trend-driven "
           "than Zudio, but more accessible than global fast-fashion labels. India is home to "
           "<b>380 million Gen Z consumers</b> — the largest youth cohort in the world. Their direct "
           "spending power is estimated at $250 billion annually, with 47% going to fashion & lifestyle. "
           "Myntra's Gen Z base alone doubled to 16 million users in a single year."),
    Paragraph("Problem 1 — Choice Overload & Decision Paralysis", sH2),
    body_p("A typical Burnt Toast collection has hundreds of products. When a 19-year-old opens the site "
           "before a college fest, they are not looking for a catalogue — they are looking for a <i>look</i>. "
           "The result? They browse for 15–20 minutes, feel overwhelmed, and either leave without buying "
           "or add one safe item they've seen before. This is <b>decision fatigue</b> — one of the biggest "
           "silent killers of online fashion conversion."),
    Paragraph("Problem 2 — No Personalisation, No Context", sH2),
    body_p("Today's website treats everyone the same. There is zero recognition of occasion, body type, "
           "style preference, or budget. Research confirms: <b>50% of fashion purchases are driven by "
           "personalisation</b>. Brands that personalise see conversion rates increase by up to 40%."),
    Paragraph("Problem 3 — The Gap Between Discovery & Purchase", sH2),
    body_p("Gen Z discovers fashion on Instagram, Pinterest, and YouTube — they see a full look in a Reel, "
           "get inspired, then land on an ecommerce site and cannot recreate that look. <b>76% of Gen Z "
           "fashion shoppers discover through social media.</b> The bridge between inspiration and purchase "
           "is broken. Toastie builds that bridge."),
    PageBreak(),
]

# Slide 3 — The Solution
story += slide_header_pdf(3, "THE SOLUTION — Meet Toastie, Your AI Fashion Stylist")
story += [
    Paragraph("<b><font color='#B8492C'>The answer is not more products. The answer is smarter discovery.</font></b>", sBody),
    Spacer(1, 4),
    body_p("Toastie is an AI-powered fashion stylist — a conversational, visual, intelligent experience "
           "built directly inside the Burnt Toast website. It is not a search bar. It is not a filter. "
           "It is not a FAQ bot. It is a <b>personal stylist</b>, available 24/7, for every customer."),
    Paragraph("What Toastie Does in Plain Language", sH2),
    body_p("A user opens Burnt Toast's website and simply talks to Toastie: <i>'I have a college fest "
           "this weekend. Y2K-inspired. Budget around ₹3,000.'</i> In seconds, Toastie builds a complete, "
           "curated outfit — top, bottom, footwear, bag, and jewelry — from Burnt Toast's real, live "
           "catalogue. Every item is shoppable. Every price is real."),
    body_p("Then they say: <i>'I love this, but can you show me different shoes?'</i> Toastie swaps just "
           "the footwear — keeping the rest intact — and shows three alternatives. The user is co-designing "
           "their look in real time, with an AI that understands fashion."),
    Paragraph("Three Things That Make Toastie Different", sH2),
    Paragraph("1. It Speaks Gen Z's Language", sH3),
    body_p("Maps casual language to precise fashion aesthetics: Y2K revival, boho-coastal, urban "
           "streetwear, feminine-romantic, minimal-clean, smart-casual, and more."),
    Paragraph("2. It Builds Complete Looks, Not Just Recommendations", sH3),
    body_p("Top, bottom, dress, outerwear, footwear, bag, jewelry, watch, hat — styled together. "
           "Color harmony, occasion fit, formality level, aesthetic coherence."),
    Paragraph("3. It's Built on Real Products, Not Generic Advice", sH3),
    body_p("Every outfit uses actual SKUs from Burnt Toast's live catalogue — real prices, real images, "
           "real sizes, direct link to purchase. The recommendation IS the product."),
    PageBreak(),
]

# Slide 4 — How It Works
story += slide_header_pdf(4, "HOW IT WORKS — The Full User Experience, Step by Step")
steps_pdf = [
    ("Step 1 — The Conversation Begins",
     "The customer sees Toastie's chat: 'Hey! What are we dressing for today? 🍞' They respond in natural "
     "language — occasion, mood, budget, vibe. Even 'I don't know, just give me something cute' is enough."),
    ("Step 2 — Toastie Reads the Brief",
     "Claude (Anthropic's AI) parses the message and identifies: Occasion, Aesthetic, Budget, Gender, "
     "Preferred colors. This is genuine language understanding — not keyword matching."),
    ("Step 3 — The Outfit is Built",
     "The engine runs across 280+ curated products and selects pieces based on aesthetic alignment, "
     "color harmony, formality level, boldness balance, and budget fit."),
    ("Step 4 — The Look is Revealed",
     "Customer sees a complete lookbook card with real images, names, prices, total cost, and a "
     "personality-driven vibe label like 'Sunday Brunch Slay' or 'Main Character Energy.'"),
    ("Step 5 — The Customer Customises",
     "Swap any piece, filter by subtype (only sneakers), change the whole vibe, lock what they love, "
     "set a budget, or add/remove items — all conversationally."),
    ("Step 6 — Add to Cart & Purchase",
     "One click goes to the product page. They arrive with intention — knowing exactly what they want, "
     "and wanting all of it."),
]
for title, text in steps_pdf:
    story.append(Paragraph(title, sH3))
    story.append(body_p(text))
story.append(PageBreak())

# Slide 5 — Key Features
story += slide_header_pdf(5, "KEY FEATURES — Full Capability List")
features_pdf = [
    ("Feature 1 — Complete Outfit Generation",
     "Full styled outfits from a natural language prompt. Covers every slot: tops, bottoms, dresses, "
     "outerwear, footwear, bags, JEWELLERY (necklaces, earrings, bracelets, rings, charms), watches, eyewear, hats."),
    ("Feature 2 — Aesthetic Intelligence (8 Style Profiles)",
     "Y2K Revival · Urban Streetwear · Smart Casual · Minimal Clean · Boho Coastal · Preppy Collegiate · Athleisure · Feminine Romantic"),
    ("Feature 3 — Occasion Mapping (11 Occasions)",
     "Casual · College · Brunch · Date Night · Party · Festival · Beach · Travel · Work · Active · Wedding"),
    ("Feature 4 — Real-Time Slot Swapping",
     "Any individual piece can be swapped without disrupting the rest of the outfit."),
    ("Feature 5 — Footwear Subtype Filter",
     "Sneakers, sandals, loafers, heels, boots, ballerinas, mary janes, mules, platforms — when asked, shows only that type. No mixing."),
    ("Feature 6 — Multi-Outfit Variations",
     "Generates 3 complete outfit options at once, each varying on a specific slot."),
    ("Feature 7 — Budget Intelligence",
     "Works within price constraints across all slots combined."),
    ("Feature 8 — Product Memory & Rejection",
     "Never re-suggests rejected products within a session — gets smarter as the conversation continues."),
    ("Feature 9 — Claude Vision (Image-Based Inspiration)",
     "Upload a photo from Instagram or Pinterest — Toastie recreates the look using Burnt Toast's real catalogue."),
    ("Feature 10 — Brand-Voice Copy Generation",
     "Every outfit gets a branded, Gen Z-native caption in Burnt Toast's voice."),
]
for title, text in features_pdf:
    story.append(Paragraph(title, sH3))
    story.append(body_p(text))
story.append(PageBreak())

# Slide 6 — Gen Z Opportunity
story += slide_header_pdf(6, "THE GEN Z OPPORTUNITY — Why Now. Why This Audience.")
story += [Paragraph("India's Gen Z — The Numbers That Matter", sH2)]
for s in [
    "<b>380 million Gen Z consumers</b> in India — largest Gen Z population of any country in the world",
    "<b>$250 billion</b> in direct annual spending; 47% goes to fashion and lifestyle",
    "By 2035, Gen Z projected to drive <b>$2 trillion</b> in total consumer spending in India",
    "<b>44%</b> of all smartphone purchases in India in 2024 came from Gen Z",
    "Myntra's Gen Z base doubled to <b>16 million users</b> in a single year",
]:
    story.append(bul(s))
story += [Paragraph("How Gen Z Shops", sH2)]
for s in [
    "<b>76%</b> discover fashion through social media — Instagram, YouTube Shorts, Pinterest",
    "<b>39%</b> of Indian Gen Z consumers buy fashion after watching Instagram Reels",
    "<b>75%</b> of Gen Z digital commerce happens on smartphones",
    "<b>50%</b> of fashion purchases driven by personalisation",
    "<b>61%</b> actively used AI tools to help with a purchase in the past year",
    "<b>75%</b> are interested in using AI during the shopping process — highest of any generation",
]:
    story.append(bul(s))
story += [
    Paragraph("The Competitive Window", sH2),
    Paragraph("<font color='#B8492C'><b>Right now, no Indian fashion brand at Burnt Toast's price point "
              "and positioning has an AI stylist. Not Zudio. Not H&amp;M India. Not Bershka. This is a "
              "first-mover window. The brand that builds the most trusted AI styling relationship with "
              "Indian Gen Z first will capture loyalty that is extremely hard to win back later.</b></font>", sBody),
    PageBreak(),
]

# Slide 7 — Business Impact: Engagement
story += slide_header_pdf(7, "BUSINESS IMPACT — Engagement")
for title, text in [
    ("Metric 1 — Time on Site",
     "Today, a typical fashion ecommerce session in India lasts 3–5 minutes. A Toastie session has "
     "narrative momentum. AI-fashion integrations globally show <b>35–60% increase in session time</b>. "
     "Every additional minute increases the statistical probability of purchase."),
    ("Metric 2 — Return Visits & Retention",
     "Toastie creates a reason to return no static catalogue can provide. Brands using AI personalisation "
     "report <b>40–60% improvement in repeat visit rates</b> among engaged users."),
    ("Metric 3 — Depth of Catalogue Engagement",
     "Without Toastie: customers discover 10–20 products per session. With Toastie: the engine surfaces "
     "products from across the entire 280+ item catalogue based on aesthetic and occasion. "
     "<b>Hidden inventory becomes discovered opportunity.</b>"),
    ("Metric 4 — Net Promoter Effect",
     "When Toastie gives someone a look they love, they screenshot it, share it on Instagram Stories, "
     "and tag Burnt Toast. <b>Organic, user-generated marketing that no ad spend can buy.</b>"),
]:
    story.append(Paragraph(title, sH2))
    story.append(Paragraph(text, sBody))
story.append(PageBreak())

# Slide 8 — Conversion & Revenue
story += slide_header_pdf(8, "BUSINESS IMPACT — Conversion & Revenue")
for title, text in [
    ("Impact 1 — Higher Conversion Rate",
     "AI-powered personalisation in fashion ecommerce increases conversion rates by <b>30–40%</b>. "
     "Visual search and AI recommendation tools show <b>27% higher conversion</b> in controlled studies. "
     "Toastie resolves the confidence problem — customers buy with intent, not hope."),
    ("Impact 2 — Larger Average Basket Size",
     "<b>Without Toastie:</b> one item, avg ₹1,200–1,500. <b>With Toastie:</b> complete outfit of 3–5 pieces, "
     "avg ₹2,500–3,500. Kering's Luce AI app reported <b>15–20% increase in average order value</b>. "
     "For Burnt Toast, every ₹500 increase across 1,000 monthly transactions = <b>₹6 lakh additional "
     "monthly revenue with zero extra acquisition cost.</b>"),
    ("Impact 3 — Reduced Return Rates",
     "When customers buy pieces that were styled together, they buy with confidence and context. "
     "Lower return rates mean lower logistics cost and higher net margin per order."),
    ("Impact 4 — New Customer Acquisition Through Experience",
     "Toastie is a reason to choose Burnt Toast specifically — not because the products are different, "
     "but because the journey of discovering them is better. "
     "<b><font color='#B8492C'>In the Gen Z market, experience IS the product.</font></b>"),
]:
    story.append(Paragraph(title, sH2))
    story.append(Paragraph(text, sBody))
story.append(PageBreak())

# Slide 9 — Chatbot Comparison
story += slide_header_pdf(9, "TOASTIE vs. A REGULAR CHATBOT")
story += [
    body_p("A standard chatbot is a customer service tool — reactive, scripted, with no understanding "
           "of aesthetics, no style logic, no ability to build something original. Toastie is a "
           "<b>generative, creative AI system</b>. The difference is fundamental."),
    Spacer(1, 6),
    comparison_table_pdf(),
    Spacer(1, 10),
]
for s in [
    "A chatbot saves time.  <b>Toastie creates desire.</b>  In fashion, desire drives purchase.",
    "A chatbot is a cost reduction tool.  <b>Toastie is a revenue generation engine.</b>",
    "A chatbot is invisible when it works.  <b>Toastie becomes a brand asset.</b>",
]:
    story.append(bul(s))
story.append(PageBreak())

# Slide 10 — Roadmap
story += slide_header_pdf(10, "FUTURE ROADMAP — 4 Phases of Growth")
roadmap = [
    ("#B8492C", "PHASE 1 — APPROVE & REFINE  (Current: Q2–Q3 2026)",
     ["Core AI stylist built, tested, and fully functional",
      "280+ product catalogue integrated across all categories",
      "Goal: panel approval, technical validation, real user testing"]),
    ("#748B6A", "PHASE 2 — LAUNCH & LEARN  (Q4 2026 – Q1 2027)",
     ["Toastie goes live on burnt-toast.com for all users",
      "Performance tracking: conversion, basket size, session time, returns",
      "New: Wishlist integration, 'Shop the Look' pages, WhatsApp/Instagram DM integration"]),
    ("#C9962E", "PHASE 3 — SCALE & PERSONALISE  (Q2–Q3 2027)",
     ["User profiles — Toastie remembers your style across sessions",
      "Size intelligence, loyalty integration, style evolution tracking",
      "In-store kiosk mode for Burnt Toast physical stores"]),
    ("#1A1A1A", "PHASE 4 — MONETISE & EXPAND  (2028+)",
     ["Trend prediction engine — proprietary signal for buying team",
      "Cross-brand expansion to Zudio, Westside, Tata CLiQ",
      "White-label licensing — AI stylist as a platform product"]),
]
for hex_col, title, items in roadmap:
    story.append(Paragraph(f"<font color='{hex_col}'><b>{title}</b></font>", sBody))
    for item in items:
        story.append(bul(item))
    story.append(Spacer(1, 4))
story.append(PageBreak())

# Slide 11 — Market Opportunity
story += slide_header_pdf(11, "MARKET OPPORTUNITY — The Size of the Prize")
story += [Paragraph("The Indian Fashion Market — Today", sH2)]
for s in [
    "India clothing & apparel market: <b>$116.64 billion in 2025</b>, growing to $177.70B by 2035",
    "Online fashion ecommerce: <b>$21.6B in 2025 → $98.45B by 2032</b>  (24.2% CAGR)",
    "Fashion ecommerce share growing from 15% (2023) to forecasted 25% by 2030",
]:
    story.append(bul(s))
story += [Paragraph("Gen Z Fashion Segment", sH2)]
for s in [
    "<b>380 million</b> Gen Z consumers in India",
    "<b>$250B</b> direct annual spending, 47% on fashion/lifestyle",
    "Projected to drive <b>$2 trillion</b> in total spending by 2035",
]:
    story.append(bul(s))
story += [Paragraph("AI Fashion Commerce — Global Context", sH2)]
for s in [
    "Generative AI in fashion market expected to reach <b>$4.4 billion by 2028</b>",
    "<b>97%</b> of commerce organisations now have AI plans in place",
    "AI-powered experiences deliver <b>30–40% conversion lift</b> and <b>15–25% basket size increase</b>",
]:
    story.append(bul(s))
story += [Paragraph("The Trent Advantage", sH2)]
for s in [
    "Trent FY25: <b>₹17,353 Cr revenue (+37%)</b>, 1,043 stores across 242 cities",
    "Zudio: 765 stores in 235 cities, <b>crossed $1 billion in sales</b> in FY25",
    "Physical Burnt Toast stores in Bengaluru, Pune, Surat, and Thane — ready for Toastie rollout",
]:
    story.append(bul(s))
story.append(PageBreak())

# Slide 12 — Customer Journey
story += slide_header_pdf(12, "CUSTOMER EXPERIENCE JOURNEY — Before vs. After Toastie")
story += [
    Paragraph("<b>Meet Priya. She's 20. She's in Mumbai. She has a college farewell party in three days.</b>", sBody),
    Paragraph("BEFORE TOASTIE", sH2),
    body_p("Priya opens Burnt Toast's website. Types 'dresses' in the search bar. 47 results. She scrolls. "
           "Some too formal, some too casual. She opens three tabs, reads descriptions, closes two. Adds "
           "one item to cart. Checks Instagram. Sees a perfect Y2K look in a Reel. Goes back to the site. "
           "Can't find it. Leaves. Buys from Myntra instead."),
    Paragraph("<font color='#B8492C'><b>Burnt Toast lost a customer. Not because the product wasn't there. "
              "Because the discovery experience failed her.</b></font>", sBody),
    Paragraph("AFTER TOASTIE", sH2),
    body_p("Priya opens Burnt Toast. Toastie: <i>'Hey! What are we dressing for? 🍞'</i> She types: "
           "<i>'College farewell. Want to look iconic. Y2K-ish. Budget ₹3,500.'</i>"),
    body_p("In 4 seconds, Toastie shows a complete look: Rust-orange co-ord set (₹1,299) · Platform "
           "sneakers (₹1,199) · Mini chain bag (₹599) · Hoop earrings + chunky bracelet (₹399) · "
           "Total: ₹3,496 ✓"),
    body_p("Label: <i>'Farewell Queen Energy.'</i>  Caption: <i>'Low-rise, high vibe. You're not leaving "
           "college — you're leaving a legacy.'</i>"),
    body_p("She wants different shoes. 'Can I see heels?' Toastie swaps just the footwear. She picks "
           "strappy kitten heels. Screenshots the look. Adds all 5 items to cart."),
    Paragraph("<font color='#748B6A'><b>₹3,496 average order. 5 products. 8-minute session. Zero ad spend.</b></font>", sBody),
    Paragraph("<font color='#B8492C'><b>Burnt Toast didn't just make a sale. It made a loyal customer.</b></font>", sBody),
    PageBreak(),
]

# Slide 13 — Development Status
story += slide_header_pdf(13, "DEVELOPMENT STATUS — What's Built, What's Next, What We Need")
story += [Paragraph("What Has Been Built (Fully Functional Now)", sH2)]
for s in [
    "✅  Conversational AI interface — Natural language, Claude-powered, Burnt Toast brand UI",
    "✅  Outfit generation engine — 280+ curated products across all categories",
    "✅  8 aesthetic profiles + 11 occasion types",
    "✅  Real-time slot swapping — any piece in any outfit",
    "✅  Footwear subtype filter — sneakers, sandals, loafers, heels, boots...",
    "✅  Budget intelligence, multi-outfit variations, product rejection memory",
    "✅  Vision AI — recreate a look from a user-uploaded photo",
    "✅  Claude-generated brand-voice copy for every outfit",
]:
    story.append(bul(s))
story += [Paragraph("Technology Stack", sH2)]
for s in [
    "Frontend: Next.js 16, React 19, TypeScript",
    "AI Layer: Anthropic Claude (Sonnet) — language understanding + image vision",
    "Outfit Engine: Custom TypeScript — built and maintained in-house",
    "Deployment: Vercel-ready — can go live within days of approval",
]:
    story.append(bul(s))
story += [Paragraph("What We Need to Proceed", sH2)]
for s in [
    "Panel Approval — integrate Toastie into the live Burnt Toast website",
    "Live Product Feed Access — Shopify API for real-time inventory sync",
    "Phase 2 Budget — 3-month engineering runway for analytics, wishlist, social sharing",
]:
    story.append(bul(s))
story.append(Paragraph("<font color='#B8492C'><b>Timeline from approval to live: 6–8 weeks.</b></font>", sBody))
story.append(PageBreak())

# Slide 14 — The Ask
story += slide_header_pdf(14, "THE ASK — What We Are Asking the TATA Panel For")
story += [
    Paragraph("<font color='#B8492C'><b>What Toastie Is — One Line:</b></font>", sBody),
    Paragraph("<i>An AI-powered personal stylist built inside Burnt Toast that turns a shopper's vibe "
              "into a complete, shoppable look in seconds — driving higher conversion, larger baskets, "
              "and deeper brand loyalty among Gen Z.</i>", sBody),
    Paragraph("Where We Stand", sH2),
]
for s in [
    "✅  Product is built and functional",
    "✅  Catalogue integrated (280+ products)",
    "✅  All core features tested and working",
    "✅  Technology stack is production-ready",
    "✅  Zero external dependencies to go live",
]:
    story.append(bul(s))
story.append(Paragraph("<font color='#B8492C'><b>All we need is the green light.</b></font>", sBody))
story.append(Paragraph("The Opportunity in Numbers", sH2))
story.append(Spacer(1, 6))
story.append(metrics_table_pdf())
story.append(Spacer(1, 10))
story += [
    Paragraph("The Closing Thought", sH2),
    body_p("Burnt Toast was built to be different. Not just another affordable fashion brand — but India's "
           "Gen Z fashion identity. Toastie is what makes that difference tangible, digital, and scalable."),
    body_p("Every other feature on a fashion website — better photography, faster shipping, easier returns "
           "— can be copied. An AI stylist that knows your vibe, builds your look, and speaks your language "
           "is a relationship. And relationships, once built, are remarkably hard to lose."),
    Paragraph("<font color='#B8492C'><b>Toastie is not a feature. It is Burnt Toast's competitive moat.</b></font>", sBody),
    Paragraph("<i><b>We are ready to build it. We just need you to say yes.</b></i>", sBody),
    Spacer(1, 1*cm),
    HR(),
    Spacer(1, 0.3*cm),
    Paragraph('<i>"Turning Scroll into Style.  Style into Sales."</i>',
              S("cq", fontSize=13, textColor=C_ACCENT, alignment=TA_CENTER, fontName="Helvetica-Oblique")),
    Paragraph("— Toastie, by Burnt Toast × TATA",
              S("cq2", fontSize=11, textColor=C_INK, alignment=TA_CENTER)),
    Spacer(1, 0.5*cm),
    Paragraph("This document is confidential and prepared for internal TATA panel review only. May 2026.",
              S("conf", fontSize=8, textColor=C_GREY, alignment=TA_CENTER)),
]

doc_pdf.build(story)
print(f"✅  Saved: {pdf_path}")
print("\n🎉  Both files are ready:")
print(f"    📄  {docx_path}  (Word Document)")
print(f"    📕  {pdf_path}  (PDF)")
